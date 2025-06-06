import whisper
from fpdf import FPDF
from docx import Document as WordDocument
from docx.shared import Pt
from dotenv import load_dotenv
from langchain.docstore.document import Document

import os
os.environ["PATH"] += os.pathsep + r"C:\Program Files\ffmpeg-7.1.1-essentials_build\bin"



def get_audio_gpt_transcript(file):
    """
    音声ファイルを読み込み、テキストに変換
    """
    model = whisper.load_model("base")
    print(f'Transcribing {file}')
    result = model.transcribe(file, language="ja")
    return Document(page_content=result["text"])

         
# 音声ファイルのリストを取得
file_list = ["14_aud.m4a", "16_aud.m4a"]

class PDF(FPDF):
    def header(self):
        self.set_font("MSGothic", size=12)
        self.cell(0, 10, "Transcription", ln=True, align="C")
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font("MSGothic", size=8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

# PDFに書き込み
#pdf = PDF()
#pdf.add_font('MSGothic', '', r"C:\phont\static\NotoSansJP-Regular.ttf", uni=True)
#pdf.add_page()
#pdf.set_auto_page_break(auto=True, margin=15)
#pdf.set_font('MSGothic', size=12)

for file in file_list:
    doc = get_audio_gpt_transcript(file)
    print("get_trans")
    word_doc = WordDocument()
    word_doc.add_heading('Transcription', level=1)

    word_doc.add_paragraph("[1]", style='Heading 2')
    para = word_doc.add_paragraph(doc.page_content)
    para.style.font.size = Pt(12)

    filename = file.replace(".m4a", "")
    word_doc.save(f"{filename}.docx")
    print("get_docx")

# transcripts は Document オブジェクトのリスト
#for doc in transcripts:
    #i = 0
    #pdf.multi_cell(0, 10, f"[{i+1}] {doc.metadata.get('source', 'N/A')}\n{doc.page_content}\n")
    #pdf.ln(5)
    #i += 1

# 保存
#pdf.output("transcripts.pdf")