import time
import fitz  # PyMuPDF
from docx import Document as WordReader
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from huggingface_hub import snapshot_download
import os

model_name = "intfloat/multilingual-e5-base"

download_path = snapshot_download(
    repo_id=model_name,
    local_dir=f"C:/Work/gemini_AI/path_to_model/{model_name}",
    local_dir_use_symlinks=False
    )


def get_lecture_title(file_path):
    """
    PDFファイルから授業のタイトルを取得する関数
    """
    if file_path.endswith(".pdf"):
        n = file_path.replace(".pdf", "")
    elif file_path.endswith(".docx"):
        n = file_path.replace(".docx", "")
    # もしnにaudが含まれていれば、授業のタイトルを取得しない
    if "aud" in n:
        n = n.replace("_aud", "")
        return f"第{n}回講義録音"
    else:
        return f"第{n}回講義資料"

# pdfファイルを読み込み、検索可能なナレッジベースを作成する関数
def create_document_base(file_path1, file_path2):
    if file_path2 is None:
        print("audio data is None")
        raw_text2 = ""
    else:
        start = time.time()
        doc = WordReader(file_path2)
        full_text = []
        for para in doc.paragraphs:
            raw_text_2 = para.text.strip()
            if raw_text_2:  # 空でない段落のみを取得
                full_text.append(raw_text_2)
        raw_text2 = "\n".join(full_text)  # 全ての段落を結合
        docx_time = time.time()
        print("get_docx_text")
        print(f"word読み込み時間: {start - docx_time:.2f}秒")
    
    if file_path1 is None:
        print("pdf data is None")
        raw_text1 = ""
    else:
        start_pdf = time.time()
        with fitz.open(file_path1) as doc1:
            raw_text1 = "\n".join(page.get_text() for page in doc1) # PDFのテキストを取得
        print("get_pdf_text")
        pdf = time.time()
        print(f"PDF読み込み時間: {pdf - start_pdf:.2f}秒")

    # テキストをチャンクに分割
    # chunk_sizeはチャンクのサイズ、chunk_overlapはオーバーラップする文字数
    str_chunk = time.time()
    text_pdf = CharacterTextSplitter(
        separator="\n",
        chunk_size=100,
        chunk_overlap=20
    ).split_text(raw_text1)
    text_docx = CharacterTextSplitter(
        separator="\n",
        chunk_size=50,
        chunk_overlap=0
    ).split_text(raw_text2)
    print("split_text")
    #print(text_docx)
    chunk = time.time()
    print(f"テキスト分割時間: {chunk - str_chunk:.2f}秒")

    # テキストをDocumentオブジェクトに変換
    # Documentオブジェクトは、page_contentとmetadataを持つ
    #metadataは、参考情報で、どのファイルから取得したかを示す
    docs_t = [Document(page_content=t, metadata = {"source" : get_lecture_title(file_path1)}) for t in text_pdf]
    docs_d = [Document(page_content=t, metadata = {"source" : get_lecture_title(file_path2)}) for t in text_docx]
    all_docs = docs_t + docs_d
    print("get_document")
    end = time.time()
    print(f"ドキュメント作成時間: {end - chunk:.2f}秒")
    return all_docs
    # HuggingFaceEmbeddings でベクトル化
    # FAISSは、ベクトルストアの一種で、ベクトル検索を行うためのライブラリ

document_base = []
# 授業のクラス名のリスト
class_list = ["1","1_1","2","3","3_1","4","4_1", "5", "6", "7", "8", "9"]

for cla in class_list:
    # PDFファイルとWordファイルのパスを指定
    if not os.path.exists(f"{cla}.pdf"):
        pdf_file = None
    else:
        pdf_file = f"{cla}.pdf"
    if not os.path.exists(f"{cla}_aud.docx"):
        docx_file = None
    else:
        docx_file = f"{cla}_aud.docx"
    #if not os.path.exists(pdf_file) or not os.path.exists(docx_file):
        #st.error(f"{cla}のファイルが見つかりません。")
        #continue
    # ナレッジベースを作成
    document_base += create_document_base(pdf_file, docx_file)

know_st = time.time()
#knowledge_base = FAISS.from_documents(document_base, HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-MiniLM-L6-v2"))
knowledge_base = FAISS.from_documents(document_base, HuggingFaceEmbeddings(model_name=download_path))
knowledge_base.save_local("faiss_index/")
print("get_knowledge_base")
know_end = time.time()
print(f"ナレッジベース作成時間: {know_end - know_st:.2f}秒")