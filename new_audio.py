import librosa
import torch
from transformers import WhisperProcessor, WhisperForConditionalGeneration
import soundfile as sf
from docx import Document as WordDocument
from docx.shared import Pt
from langchain.docstore.document import Document
import math
from pydub.silence import split_on_silence

# 設定
model_id = "kotoba-tech/kotoba-whisper-v2.0"
device = "cuda" if torch.cuda.is_available() else "cpu"

# プロセッサとモデルのロード
processor = WhisperProcessor.from_pretrained(model_id)
model = WhisperForConditionalGeneration.from_pretrained(model_id).to(device)
model.config.attn_implementation = "flash_attention_2"

def convert_m4a_to_wav(input_path, output_path):
  """
  M4AファイルをWAVファイルに変換する関数

  Args:
    input_path: 変換するM4Aファイルのパス
    output_path: 出力WAVファイルのパス
  """
  try:
    y, sr = librosa.load(input_path, sr=None)
    sf.write(output_path, y, sr)
    print(f"変換完了: {input_path} -> {output_path}")
  except Exception as e:
    print(f"変換エラー: {e}")

def split_aud(fi, sr = 16000, chunk_sec=30):
    chunk_size = sr*chunk_sec
    total_chunks = math.ceil(len(fi) / chunk_size)
    return [fi[i*chunk_size:(i+1)*chunk_size] for i in range(total_chunks)]


def get_string(chunk):

    # 音声の読み込みとリサンプリング
    #audio_input, _ = librosa.load(aud_file, sr=16000)  # sr=16000で自動リサンプリング
    #audio_input, sample_rate = sf.read(aud_chunk)

    # 音声データの前処理
    #inputs = processor(audio_input, sampling_rate=16000, return_tensors="pt")
    input_features = processor(chunk, sampling_rate=16000, return_tensors="pt").input_features.to(device)

    # 推論の実行
    generated_ids = model.generate(input_features)
    transcription = processor.batch_decode(generated_ids, skip_special_tokens=True)
    #print(transcription)
    return transcription

file_list = ["7_aud.m4a"]

for file in file_list:
    filename = file.replace(".wav", "")
    convert_m4a_to_wav(file,f"{filename}.wav")
    audio_input, _ = librosa.load(file, sr=16000)
    chunk_list = split_aud(audio_input)
    print(chunk_list)
    i = 1
    word_doc = WordDocument()
    word_doc.add_heading(f'{file}書き起こし', level=1)
    for chunk in chunk_list:
        doc = get_string(chunk)
        print("get_trans")
        word_doc.add_paragraph(f"[{i}]", style='Heading 2')
        para = word_doc.add_paragraph(doc[0])
        para.style.font.size = Pt(12)
        i += 1
    word_doc.save(f"{filename}.docx")
    print("get_docx")

