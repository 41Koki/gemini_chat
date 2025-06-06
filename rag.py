import os
from dotenv import load_dotenv
import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.document_loaders import AssemblyAIAudioTranscriptLoader
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
from langchain.text_splitter import CharacterTextSplitter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document as WordReader
from langchain.docstore.document import Document
from langchain.docstore.document import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
import fitz
from fpdf import FPDF
import time


# 同一フォルダにある.envファイルを読み込む
load_dotenv()
# モデルはGemini-1.5-flashを指定
llm = ChatGoogleGenerativeAI(model='gemini-1.5-flash')

def get_lecture_title(file_path):
    """
    PDFファイルから授業のタイトルを取得する関数
    """
    if file_path.endswith(".pdf"):
        n = file_path.replace(".pdf", "")
    elif file_path.endswith(".docx"):
        n = file_path.replace(".docx", "")
    # もしnにaudが含まれていれば、授業のタイトルを取得しない
    if "aud" in n:
        return f"第{n}回講義録音"
    else:
        return f"第{n}回講義資料"

# pdfファイルを読み込み、検索可能なナレッジベースを作成する関数
def create_document_base(file_path1, file_path2):
    start = time.time()
    doc = WordReader(file_path2)
    with fitz.open(file_path1) as doc1:
        raw_text1 = "\n".join(page.get_text() for page in doc1) # PDFのテキストを取得
    print("get_pdf_text")
    pdf = time.time()
    full_text = []
    for para in doc.paragraphs:
        raw_text_2 = para.text.strip()
        if raw_text_2:  # 空でない段落のみを取得
            full_text.append(raw_text_2)
    raw_text2 = "\n".join(full_text)  # 全ての段落を結合
    print("get_docx_text")
    docx = time.time()
    print(f"PDF読み込み時間: {pdf - start:.2f}秒")
    print(f"Word読み込み時間: {docx - pdf:.2f}秒")

    # テキストをチャンクに分割
    # chunk_sizeはチャンクのサイズ、chunk_overlapはオーバーラップする文字数
    text_pdf = CharacterTextSplitter(
        separator="\n",
        chunk_size=100,
        chunk_overlap=20
    ).split_text(raw_text1)
    text_docx = CharacterTextSplitter(
        separator="\n",
        chunk_size=10,
        chunk_overlap=0
    ).split_text(raw_text2)
    print("split_text")
    chunk = time.time()
    print(f"テキスト分割時間: {chunk - docx:.2f}秒")

    # テキストをDocumentオブジェクトに変換
    # Documentオブジェクトは、page_contentとmetadataを持つ
    #metadataは、参考情報で、どのファイルから取得したかを示す
    docs_t = [Document(page_content=t, metadata = {"source" : get_lecture_title(file_path1)}) for t in text_pdf]
    docs_d = [Document(page_content=t, metadata = {"source" : get_lecture_title(file_path2)}) for t in text_docx]
    all_docs = docs_t + docs_d
    print("get_document")
    end = time.time()
    print(f"ドキュメント作成時間: {end - chunk:.2f}秒")
    return all_docs
    # HuggingFaceEmbeddings でベクトル化
    # FAISSは、ベクトルストアの一種で、ベクトル検索を行うためのライブラリ

document_base = []
# 授業のクラス名のリスト
class_list = ["14","16"]

for cla in class_list:
    # PDFファイルとWordファイルのパスを指定
    pdf_file = f"{cla}.pdf"
    docx_file = f"{cla}.docx"
    #if not os.path.exists(pdf_file) or not os.path.exists(docx_file):
        #st.error(f"{cla}のファイルが見つかりません。")
        #continue
    # ナレッジベースを作成
    document_base += create_document_base(pdf_file, docx_file)

know_st = time.time()
knowledge_base = FAISS.from_documents(document_base, HuggingFaceEmbeddings(model_name="all-mpnet-base-v2"))
print("get_knowledge_base")
know_end = time.time()
print(f"ナレッジベース作成時間: {know_end - know_st:.2f}秒")



# 初期メッセージ
system_message = SystemMessage(content="あなたは授業アシスタントです。\n\
                                        授業の内容に関する質問に答えることができます。\n\
                                        質問に対する答えは、参考情報をもとに生成してください。\n\
                                        参考資料で、pdfとdocxで同じ数字がファイル名に含まれるものは、同じ内容について記載されています。\n\
                                        もし、参考情報に関係する単語が全くない場合は、関係する単語が見つからなかったと回答してください。\n\
                                        もし、少しだけ関連する単語がある場合は、関連する単語を使って回答してください。\n")

                                        
# 会話履歴を格納するリスト
# のちにボタンを押したとき、会話履歴がリセットされるのを防ぐために、
# session_stateにconversation_historyを格納
# session_stateは、Streamlitのセッション状態を管理するための辞書型オブジェクト
# 最初だけ初期メッセージを追加
if "conversation_history" not in st.session_state:
    st.session_state.conversation_history = [system_message]
conversation_history = st.session_state.conversation_history

# StreamlitのUIを作成
st.title("langchain-streamlit-app")

if "messages" not in st.session_state:
    st.session_state.messages = []

# 会話履歴を画面に表示するためのリストであるmessagesを作成
# messagesは、ユーザーとAIのメッセージを格納するリスト
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# ユーザーからの入力を受け取る
# chat_inputは、ユーザーからの入力を受け取るためのStreamlitのウィジェット
prompt = st.chat_input("お困りのことがあれば教えてください。")

if prompt:
    pro_st = time.time()
    st.session_state.messages.append({"role": "user", "content": prompt})
    # ユーザーからの入力を受け取ったら、ナレッジベースから関連する情報を取得
    retriever = knowledge_base.as_retriever()
    retriever.search_type = "similarity" # 類似度検索を使用
    retriever.search_kwargs = {"k": 3} # 上位3件の情報を取得
    print("get_retriever")
    retriever_st = time.time()
    print(f"リトリーバー取得時間: {retriever_st - pro_st:.2f}秒")
    context_text = "\n\n".join([
                                f"[{doc.metadata.get('source')}]\n{doc.page_content}" 
                                for doc in retriever.invoke(prompt)]) # ユーザーからの入力に関連する情報を取得
    print("get_context_text")
    context_st = time.time()
    print(f"コンテキスト取得時間: {context_st - retriever_st:.2f}秒")
    gen_prompt = f"質問: {prompt}\n\n以下は、参考情報です。\n\n{context_text}\n" 
    context_end = time.time()
    print(f"プロンプト生成時間: {context_end - context_st:.2f}秒")
    st.session_state.information = f"以下は、参考情報です。\n{context_text}"
    conversation_history.append(HumanMessage(content=gen_prompt)) # 会話履歴にユーザーからの入力と参考情報を追加


    with st.chat_message("user"):
        st.markdown(prompt) # ユーザーからの入力を表示
    with st.chat_message("assistant"):
        response_st = time.time()
        response = llm.invoke(conversation_history) # モデルに会話履歴を渡して応答を生成
        response_end = time.time()
        print(f"応答生成時間: {response_end - response_st:.2f}秒")
        # 会話履歴にAIの応答を追加
        conversation_history.append(AIMessage(content=response.content))
        st.markdown(response.content) # AIの応答を表示
    st.session_state.messages.append({"role": "assistant", "content": response.content}) 

# 参考情報を表示するボタンを作成
# if promptの外に表示しないと、ボタンが押されても情報が表示されない
if st.button("参考情報を表示"):
    if "information" in st.session_state:
        st.markdown(st.session_state.information)
